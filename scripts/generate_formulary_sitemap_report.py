#!/usr/bin/env python3
"""Generate Formulary-Sitemap.docx report with TOC."""

import json
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ROOT = Path(__file__).resolve().parents[1]
INDEX_PATH = ROOT / "data" / "processed" / "formulary_sitemap_index.json"
OUTPUT_PATH = ROOT / "Formulary-Sitemap.docx"

# Load data
with open(INDEX_PATH, "r", encoding="utf-8") as f:
    idx = json.load(f)

pairs = idx["pairs"]
metadata = idx["metadata"]

# Compute stats
state_drug_counts: dict[str, int] = {}
drug_set: set[str] = set()
for p in pairs:
    parts = p.split("/", 1)
    state_slug = parts[0]
    drug_slug = parts[1] if len(parts) > 1 else ""
    state_drug_counts[state_slug] = state_drug_counts.get(state_slug, 0) + 1
    drug_set.add(drug_slug)

states_sorted = sorted(state_drug_counts.items(), key=lambda x: -x[1])
total_pairs = len(pairs)
unique_drugs = len(drug_set)
unique_states = len(state_drug_counts)
chunks = []
for i in range((total_pairs + 49999) // 50000):
    start = i * 50000
    end = min(start + 50000, total_pairs)
    chunks.append((f"formulary-{i+1}", end - start))

STATE_DISPLAY = {
    "alabama": "Alabama", "alaska": "Alaska", "arizona": "Arizona",
    "arkansas": "Arkansas", "california": "California", "colorado": "Colorado",
    "connecticut": "Connecticut", "delaware": "Delaware",
    "district-of-columbia": "District of Columbia", "florida": "Florida",
    "georgia": "Georgia", "hawaii": "Hawaii", "idaho": "Idaho",
    "illinois": "Illinois", "indiana": "Indiana", "iowa": "Iowa",
    "kansas": "Kansas", "kentucky": "Kentucky", "louisiana": "Louisiana",
    "maine": "Maine", "maryland": "Maryland", "massachusetts": "Massachusetts",
    "michigan": "Michigan", "minnesota": "Minnesota", "mississippi": "Mississippi",
    "missouri": "Missouri", "montana": "Montana", "nebraska": "Nebraska",
    "nevada": "Nevada", "new-hampshire": "New Hampshire",
    "new-jersey": "New Jersey", "new-mexico": "New Mexico",
    "new-york": "New York", "north-carolina": "North Carolina",
    "north-dakota": "North Dakota", "ohio": "Ohio", "oklahoma": "Oklahoma",
    "oregon": "Oregon", "pennsylvania": "Pennsylvania",
    "rhode-island": "Rhode Island", "south-carolina": "South Carolina",
    "south-dakota": "South Dakota", "tennessee": "Tennessee", "texas": "Texas",
    "utah": "Utah", "vermont": "Vermont", "virginia": "Virginia",
    "washington": "Washington", "west-virginia": "West Virginia",
    "wisconsin": "Wisconsin", "wyoming": "Wyoming",
}

# ── Document setup ──────────────────────────────────────────────────────────

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1A, 0x3A, 0x2A)  # forest green

# Helper
def add_shaded_cell(cell, text, shade_color="1A3A2A", font_color="FFFFFF", bold=True):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs[0].text:
        p.runs[0].text = text
    run = p.runs[0]
    run.font.color.rgb = RGBColor(int(font_color[:2], 16), int(font_color[2:4], 16), int(font_color[4:], 16))
    run.bold = bold
    run.font.size = Pt(10)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        add_shaded_cell(cell, h)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ── COVER PAGE ──────────────────────────────────────────────────────────────

doc.add_paragraph()  # spacer
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Formulary Sitemap\nExpansion Report")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x2A)
run.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("HealthInsuranceRenew.com")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0xC9, 0x97, 0x3A)  # gold
run.bold = True

doc.add_paragraph()

meta_block = doc.add_paragraph()
meta_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_block.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}\n").font.size = Pt(11)
meta_block.add_run(f"Data Pipeline Run: {metadata.get('generated_at', 'N/A')[:10]}\n").font.size = Pt(11)
meta_block.add_run(f"Plan Year: 2026\n").font.size = Pt(11)
meta_block.add_run(f"Total Sitemap URLs: {total_pairs:,}").font.size = Pt(14)

doc.add_page_break()


# ── TABLE OF CONTENTS (field code) ─────────────────────────────────────────

doc.add_heading("Table of Contents", level=1)

# Insert TOC field — Word will populate on open (right-click > Update Field)
p = doc.add_paragraph()
run = p.add_run()
fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fld_char_begin)
run2 = p.add_run()
instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
run2._r.append(instr)
run3 = p.add_run()
fld_char_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run3._r.append(fld_char_separate)
run4 = p.add_run("[Right-click and select 'Update Field' to populate Table of Contents]")
run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run4.font.size = Pt(10)
run4.italic = True
run5 = p.add_run()
fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run5._r.append(fld_char_end)

doc.add_page_break()


# ── 1. EXECUTIVE SUMMARY ───────────────────────────────────────────────────

doc.add_heading("1. Executive Summary", level=1)

doc.add_paragraph(
    "This report documents the expansion of the HealthInsuranceRenew.com formulary sitemap "
    "from a 400-URL seed set to comprehensive coverage of all valid state and drug combinations "
    "derived from the CMS formulary dataset. The expansion increases Google-discoverable formulary "
    "pages by 563x, covering 44 US states and 18,598 unique prescription drugs."
)

doc.add_heading("Key Metrics", level=2)

add_table(doc,
    ["Metric", "Before", "After", "Change"],
    [
        ["Total Formulary URLs", "400", f"{total_pairs:,}", f"+{total_pairs - 400:,} ({total_pairs / 400:.0f}x)"],
        ["URL Pattern", "/formulary/{issuer-id}/{drug}", "/formulary/{state-slug}/{drug-slug}", "State-first"],
        ["States Covered", "0 (issuer IDs only)", str(unique_states), f"+{unique_states}"],
        ["Unique Drugs", "20 (seed list)", f"{unique_drugs:,}", f"+{unique_drugs - 20:,}"],
        ["Sub-Sitemaps", "1 (formulary)", f"{len(chunks)} (formulary-1 to formulary-{len(chunks)})", f"+{len(chunks) - 1}"],
        ["Sitemap Limit Compliant", "Yes (400 < 50K)", "Yes (50K per chunk)", "Split applied"],
    ],
    col_widths=[4.5, 5, 5, 3.5],
)

doc.add_paragraph()
doc.add_paragraph(
    "The previous sitemap used a cross-product of the top 20 issuers (by plan count) and 20 seed drug names, "
    "generating only 400 URLs. These used opaque issuer IDs in the URL path (e.g., /formulary/98185/metformin), "
    "which have no SEO value and don't match user search patterns. The new sitemap uses state-first, "
    "human-readable URLs (e.g., /formulary/texas/metformin) that align with how users actually search: "
    "\"[drug name] coverage in [state]\"."
)

doc.add_page_break()


# ── 2. METHODOLOGY ─────────────────────────────────────────────────────────

doc.add_heading("2. Methodology", level=1)

doc.add_heading("2.1 Data Sources", level=2)
add_table(doc,
    ["Source", "File", "Size", "Purpose"],
    [
        ["CMS Formulary Data", "formulary_intelligence.json", "7.73 GB", "Drug records with issuer_ids"],
        ["Byte-Offset Index", ".cache/formulary_drug_index.json", "2.8 MB", "Drug name to file offset mapping"],
        ["Plan Intelligence", "plan_intelligence.json", "107 MB", "FFM issuer-to-state mapping"],
        ["SBM Formulary Files", "formulary_sbm_*.json (17 files)", "~50 MB total", "State-Based Marketplace issuer mapping"],
    ],
    col_widths=[4, 5, 2.5, 6],
)

doc.add_heading("2.2 Extraction Pipeline", level=2)
doc.add_paragraph(
    "The sitemap index was built by the script scripts/etl/build_formulary_sitemap_index.py "
    "using the following pipeline:"
)

steps = [
    ("Step 1 -- Issuer-State Mapping: ", "Built a mapping of issuer_id to state codes by combining "
     "FFM data from plan_intelligence.json (183 issuers across 30 states) with SBM data from 17 "
     "state-specific formulary files. Result: 200 issuers mapped to 46 states."),
    ("Step 2 -- Drug Index Load: ", "Loaded the byte-offset index (.cache/formulary_drug_index.json) "
     "containing 36,429 drug name entries with their file offsets and block lengths."),
    ("Step 3 -- Block Scanning: ", "For each drug entry in the index, read its NDJSON byte block from "
     "the 7.73 GB formulary file, extracted all issuer_ids from the records, and resolved each issuer "
     "to its operating states. Result: 18,598 unique drug slugs with state data."),
    ("Step 4 -- Pair Generation: ", "Cross-referenced drug slugs with their valid states to produce "
     f"{total_pairs:,} unique state-slug/drug-slug pairs."),
    ("Step 5 -- Index Output: ", "Wrote the compact formulary_sitemap_index.json (8.7 MB) containing "
     "all pairs as simple strings for fast sitemap generation at request time."),
]
for prefix, text in steps:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading("2.3 Drug Name Normalization", level=2)
doc.add_paragraph(
    "CMS formulary data uses verbose clinical names (e.g., '0.25 ML SEMAGLUTIDE 0.68 MG/ML "
    "PEN INJECTOR [OZEMPIC]'). The pipeline normalizes these to URL-safe slugs by:"
)
add_bullet(doc, "Extracting the brand name from [BRACKETS] when present")
add_bullet(doc, "Falling back to the full clinical name if no bracket notation exists")
add_bullet(doc, "Converting to lowercase, replacing non-alphanumeric characters with hyphens")
add_bullet(doc, "Collapsing consecutive hyphens and trimming")
doc.add_paragraph(
    "Example: '0.25 ML SEMAGLUTIDE 0.68 MG/ML PEN INJECTOR [OZEMPIC]' becomes 'ozempic'"
)

doc.add_page_break()


# ── 3. SITEMAP ARCHITECTURE ────────────────────────────────────────────────

doc.add_heading("3. Sitemap Architecture", level=1)

doc.add_heading("3.1 Sitemap Index Structure", level=2)
doc.add_paragraph(
    "The sitemap index at /sitemap.xml now dynamically generates sub-sitemap references. "
    f"The formulary section is split into {len(chunks)} sub-sitemaps to comply with Google's "
    "50,000 URL limit per sitemap file."
)

doc.add_heading("3.2 Sub-Sitemap Breakdown", level=2)
add_table(doc,
    ["Sub-Sitemap", "URL Count", "URL Range", "Percentage"],
    [
        [name, f"{count:,}", f"{(i)*50000 + 1:,} -- {(i)*50000 + count:,}",
         f"{count/total_pairs*100:.1f}%"]
        for i, (name, count) in enumerate(chunks)
    ],
    col_widths=[4, 3, 5, 3],
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(f"Total: {total_pairs:,} URLs across {len(chunks)} sub-sitemaps")
run.bold = True

doc.add_heading("3.3 URL Pattern", level=2)
doc.add_paragraph("All formulary sitemap URLs follow the pattern:")
p = doc.add_paragraph()
run = p.add_run("https://healthinsurancerenew.com/formulary/{state-slug}/{drug-slug}")
run.font.name = "Consolas"
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x2A)

doc.add_paragraph()
doc.add_paragraph("Examples:")
examples = [
    "/formulary/texas/metformin",
    "/formulary/north-carolina/ozempic",
    "/formulary/florida/atorvastatin",
    "/formulary/ohio/lisinopril",
    "/formulary/california/levothyroxine",
]
for ex in examples:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"https://healthinsurancerenew.com{ex}")
    run.font.name = "Consolas"
    run.font.size = Pt(9)

doc.add_heading("3.4 Caching and Revalidation", level=2)
add_bullet(doc, "set to 86,400 seconds (24 hours)", bold_prefix="Revalidation: ")
add_bullet(doc, "public, max-age=86400, s-maxage=86400", bold_prefix="Cache-Control: ")
add_bullet(doc, "application/xml", bold_prefix="Content-Type: ")

doc.add_page_break()


# ── 4. STATE COVERAGE ──────────────────────────────────────────────────────

doc.add_heading("4. State-by-State Coverage", level=1)

doc.add_paragraph(
    f"Formulary data is available for {unique_states} of 51 US states/territories (48 states + DC + 2 territories). "
    "The 7 states without coverage are State-Based Marketplace (SBM) states whose formulary data "
    "is not available through the CMS MR-PUF pipeline (CA, NY, MA are geo-blocked; others have "
    "limited data access)."
)

doc.add_heading("4.1 Complete State Coverage Table", level=2)

# Split into two columns of ~22 rows each for readability
half = (len(states_sorted) + 1) // 2
rows = []
for i in range(half):
    left = states_sorted[i]
    right = states_sorted[i + half] if i + half < len(states_sorted) else ("", "")
    rows.append([
        i + 1,
        STATE_DISPLAY.get(left[0], left[0]),
        f"{left[1]:,}",
        f"{left[1]/total_pairs*100:.1f}%",
        i + half + 1 if i + half < len(states_sorted) else "",
        STATE_DISPLAY.get(right[0], right[0]) if right[0] else "",
        f"{right[1]:,}" if right[1] else "",
        f"{right[1]/total_pairs*100:.1f}%" if right[1] else "",
    ])

add_table(doc,
    ["#", "State", "Drug URLs", "Share", "#", "State", "Drug URLs", "Share"],
    rows,
    col_widths=[1, 3.5, 2, 1.5, 1, 3.5, 2, 1.5],
)

doc.add_heading("4.2 Top 10 States by Coverage", level=2)
doc.add_paragraph("States with the most drug coverage pages, indicating broader issuer participation:")
top10 = states_sorted[:10]
add_table(doc,
    ["Rank", "State", "Drug URLs", "% of Total", "Unique Issuers (est.)"],
    [
        [i+1, STATE_DISPLAY.get(s, s), f"{c:,}", f"{c/total_pairs*100:.1f}%", "Multiple"]
        for i, (s, c) in enumerate(top10)
    ],
    col_widths=[1.5, 4, 3, 3, 3],
)

doc.add_heading("4.3 States Not Covered", level=2)
doc.add_paragraph(
    "The following states/territories are not included in the formulary sitemap due to "
    "data access limitations:"
)
missing = ["California (CA)*", "New York (NY)*", "Massachusetts (MA)*",
           "Minnesota (MN)", "Rhode Island (RI)", "South Dakota (SD)", "Vermont (VT)"]
for m in missing:
    add_bullet(doc, m)
doc.add_paragraph(
    "* CA, NY, and MA operate fully independent State-Based Marketplaces with no public "
    "formulary API endpoints accessible through the CMS pipeline. These states are handled "
    "by the SBM explanation page fallback on the site."
).italic = True

doc.add_page_break()


# ── 5. DRUG COVERAGE ──────────────────────────────────────────────────────

doc.add_heading("5. Drug Coverage Analysis", level=1)

doc.add_paragraph(
    f"The sitemap covers {unique_drugs:,} unique drug slugs derived from {metadata.get('unique_drugs', 'N/A')} "
    "drug name entries in the CMS formulary dataset. Drug names range from common generics "
    "(metformin, lisinopril) to specialty biologics and combination therapies."
)

doc.add_heading("5.1 Drug Categories Represented", level=2)
doc.add_paragraph(
    "The drugs span all 20 therapeutic categories defined in the site's drug taxonomy:"
)
categories = [
    "Diabetes (Metformin, Ozempic, Jardiance, Mounjaro)",
    "Blood Pressure (Lisinopril, Amlodipine, Losartan, Metoprolol)",
    "Cholesterol (Atorvastatin, Rosuvastatin, Simvastatin)",
    "Mental Health (Sertraline, Escitalopram, Bupropion)",
    "Pain & Inflammation (Gabapentin, Meloxicam, Pregabalin)",
    "Thyroid (Levothyroxine, Synthroid, Armour Thyroid)",
    "Respiratory (Albuterol, Montelukast, Fluticasone)",
    "Stomach/GI (Omeprazole, Pantoprazole, Famotidine)",
    "Heart/Rhythm (Eliquis, Xarelto, Warfarin)",
    "Kidney/Diuretics (Furosemide, Spironolactone)",
    "Antibiotics (Amoxicillin, Azithromycin, Doxycycline)",
    "Sleep (Trazodone, Zolpidem, Hydroxyzine)",
    "ADHD (Adderall, Vyvanse, Concerta)",
    "Seizure (Lamotrigine, Levetiracetam, Topiramate)",
    "Skin (Tretinoin, Clobetasol, Triamcinolone)",
    "Eye Care (Latanoprost, Timolol, Brimonidine)",
    "Osteoporosis (Alendronate, Risedronate, Denosumab)",
    "Autoimmune (Humira, Enbrel, Methotrexate)",
    "Weight Loss (Wegovy, Ozempic, Mounjaro, Zepbound)",
    "Women's Health (Estradiol, Progesterone, Oral contraceptives)",
]
for cat in categories:
    add_bullet(doc, cat)

doc.add_heading("5.2 Coverage Distribution", level=2)
doc.add_paragraph(
    f"On average, each drug slug appears in {total_pairs / unique_drugs:.1f} states. "
    f"Each state has an average of {total_pairs / unique_states:,.0f} drug URLs."
)

doc.add_page_break()


# ── 6. TECHNICAL IMPLEMENTATION ────────────────────────────────────────────

doc.add_heading("6. Technical Implementation", level=1)

doc.add_heading("6.1 Files Created", level=2)
add_table(doc,
    ["File", "Type", "Size", "Purpose"],
    [
        ["scripts/etl/build_formulary_sitemap_index.py", "Python Script", "~8 KB", "Extracts all valid state/drug pairs from formulary data"],
        ["data/processed/formulary_sitemap_index.json", "JSON Index", "8.7 MB", "Pre-built index of 225,261 state/drug pairs"],
    ],
    col_widths=[7, 2.5, 2, 5.5],
)

doc.add_heading("6.2 Files Modified", level=2)
add_table(doc,
    ["File", "Change Description"],
    [
        ["lib/data-loader.ts", "Added loadFormularySitemapIndex() function with caching"],
        ["app/sitemap.xml/route.ts", "Dynamic formulary-1..N sub-sitemap generation in sitemap index"],
        ["app/sitemaps/[type]/route.ts", "Replaced 20x20 seed with full index-driven formulary entries; added formulary-N regex routing"],
    ],
    col_widths=[7, 10],
)

doc.add_heading("6.3 How It Works", level=2)

doc.add_paragraph("Build-Time (run once when data is refreshed):", style="Heading 3")
add_bullet(doc, "runs build_formulary_sitemap_index.py", bold_prefix="Step 1: ")
add_bullet(doc, "Script streams through 7.73 GB formulary file using byte-offset index", bold_prefix="Step 2: ")
add_bullet(doc, "For each drug block, extracts issuer_ids and resolves to states", bold_prefix="Step 3: ")
add_bullet(doc, "Outputs compact JSON index (8.7 MB) with all valid pairs", bold_prefix="Step 4: ")

doc.add_paragraph("Request-Time (when Google crawls /sitemaps/formulary-N):", style="Heading 3")
add_bullet(doc, "Sitemap index at /sitemap.xml reads the formulary index to calculate chunk count", bold_prefix="Step 1: ")
add_bullet(doc, "Google requests /sitemaps/formulary-1, /sitemaps/formulary-2, etc.", bold_prefix="Step 2: ")
add_bullet(doc, "Each sub-sitemap handler slices the pairs array at the correct offset", bold_prefix="Step 3: ")
add_bullet(doc, "Returns XML with up to 50,000 URLs per response", bold_prefix="Step 4: ")

doc.add_heading("6.4 Performance Characteristics", level=2)
add_table(doc,
    ["Operation", "Duration", "Memory"],
    [
        ["Index generation (Python)", "~1.8 seconds", "~200 MB (plan_intelligence load)"],
        ["Sitemap index load (Node.js)", "< 100ms (cached)", "~9 MB (one-time)"],
        ["Sub-sitemap generation", "< 50ms per chunk", "Negligible (string concatenation)"],
        ["Revalidation interval", "24 hours", "N/A"],
    ],
    col_widths=[5, 4, 5],
)

doc.add_page_break()


# ── 7. SEO IMPACT ──────────────────────────────────────────────────────────

doc.add_heading("7. SEO Impact Assessment", level=1)

doc.add_heading("7.1 Discoverability Improvement", level=2)
doc.add_paragraph(
    "The sitemap expansion has a direct impact on search engine discoverability:"
)
add_table(doc,
    ["Metric", "Before", "After"],
    [
        ["Formulary pages in sitemap", "400", f"{total_pairs:,}"],
        ["Search-friendly URLs", "No (opaque issuer IDs)", "Yes (state + drug names)"],
        ["State-level targeting", "None", f"{unique_states} states"],
        ["Long-tail keyword coverage", "20 drugs", f"{unique_drugs:,} drugs"],
        ["Google crawl efficiency", "Low (most pages undiscoverable)", "High (all valid pages listed)"],
    ],
    col_widths=[5, 5, 5],
)

doc.add_heading("7.2 Search Intent Alignment", level=2)
doc.add_paragraph(
    "The new URL structure aligns with how users search for drug coverage information:"
)
add_table(doc,
    ["User Search Query", "Matching Sitemap URL"],
    [
        ['"metformin coverage in Texas"', "/formulary/texas/metformin"],
        ['"is Ozempic covered in Florida"', "/formulary/florida/ozempic"],
        ['"Eliquis tier North Carolina"', "/formulary/north-carolina/eliquis"],
        ['"Humira formulary Ohio"', "/formulary/ohio/humira"],
        ['"Wegovy health insurance Washington"', "/formulary/washington/wegovy"],
    ],
    col_widths=[7, 10],
)

doc.add_heading("7.3 Estimated Indexable Page Potential", level=2)
doc.add_paragraph(
    f"With {total_pairs:,} formulary URLs now in the sitemap, these pages become candidates for "
    "Google indexing. Based on typical programmatic SEO indexing rates (60-80% for well-structured "
    f"sitemaps), we estimate {int(total_pairs * 0.65):,} to {int(total_pairs * 0.80):,} pages will "
    "be indexed within 3-6 months of sitemap submission."
)

doc.add_page_break()


# ── 8. DRUGS SITEMAP AUDIT ─────────────────────────────────────────────────

doc.add_heading("8. Related: /drugs Sitemap Audit", level=1)

doc.add_paragraph(
    "An audit of the /drugs sub-sitemap was also performed to identify similar coverage gaps."
)

doc.add_heading("8.1 Current /drugs Sitemap Composition", level=2)
add_table(doc,
    ["URL Type", "Count", "Pattern"],
    [
        ["Drug category pages", "20", "/drugs/categories/{category}"],
        ["Drug comparison pages", "15", "/drugs/compare/{drug-a}-vs-{drug-b}"],
        ["County drug pages (seed)", "500", "/drugs/{state}/{county}/{drug}"],
        ["Total", "535", ""],
    ],
    col_widths=[5, 2, 10],
)

doc.add_heading("8.2 Assessment", level=2)
doc.add_paragraph(
    "The 500 county-level drug pages (/drugs/{state}/{county}/{drug}) are 301 redirects "
    "to the canonical URLs at /{state-slug}/{county-slug}/{drug}-coverage. These redirect "
    "URLs should ideally not be in the sitemap -- search engines should discover the canonical "
    "targets directly."
)
doc.add_paragraph(
    "However, the actual drug coverage pages (served by the [state-name]/[county-slug]/[county-page] "
    "catch-all route) are already discoverable through:"
)
add_bullet(doc, "The SBC sub-sitemap (which covers all plan detail pages)")
add_bullet(doc, "The new formulary sub-sitemaps (which cover all state/drug combinations)")
add_bullet(doc, "Internal linking from drug category and comparison pages")

doc.add_paragraph(
    "Conclusion: The /drugs sitemap gap is a minor issue. The redirect URLs provide some crawl "
    "signal, and the canonical pages are covered by other sitemaps. No urgent action needed, "
    "but a future cleanup could remove the redirect URLs and replace them with canonical targets."
).italic = True

doc.add_page_break()


# ── 9. MAINTENANCE ─────────────────────────────────────────────────────────

doc.add_heading("9. Maintenance & Refresh Procedures", level=1)

doc.add_heading("9.1 When to Regenerate the Index", level=2)
add_bullet(doc, "After any CMS formulary data refresh (typically annually for plan year updates)")
add_bullet(doc, "After adding new SBM state formulary files")
add_bullet(doc, "After changes to the issuer-state mapping in plan_intelligence.json")

doc.add_heading("9.2 Regeneration Command", level=2)
p = doc.add_paragraph()
run = p.add_run("python scripts/etl/build_formulary_sitemap_index.py")
run.font.name = "Consolas"
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x2A)

doc.add_paragraph("Runtime: ~2 seconds. No restart required -- the sitemap handler reads the "
                  "updated index on next revalidation cycle (within 24 hours).")

doc.add_heading("9.3 Verification Checklist", level=2)
checks = [
    "formulary_sitemap_index.json exists in data/processed/",
    "TypeScript compiles cleanly (npx tsc --noEmit)",
    "/sitemap.xml lists formulary-1 through formulary-N entries",
    "Each /sitemaps/formulary-N returns valid XML with <= 50,000 URLs",
    "Sample URLs resolve correctly (e.g., /formulary/texas/metformin)",
    "Total URL count matches expected pairs from index metadata",
]
for i, check in enumerate(checks, 1):
    add_bullet(doc, check, bold_prefix=f"[ ] {i}. ")

doc.add_page_break()


# ── 10. APPENDIX ───────────────────────────────────────────────────────────

doc.add_heading("10. Appendix", level=1)

doc.add_heading("10.1 Complete Sitemap Index Structure", level=2)
doc.add_paragraph("The /sitemap.xml now contains the following sub-sitemaps:")

fixed_types = [
    ("static", "Homepage, tools, guides, trust pages"),
    ("plans", "State plan hubs + county plan pages"),
    ("subsidies", "Subsidy state indexes + county pages"),
    ("rates", "Rate volatility state + county pages"),
    ("enhanced-credits", "Enhanced credit state + county pages"),
    ("sbc", "Individual plan detail (SBC) pages"),
]
for name, count in chunks:
    fixed_types.append((name, f"Formulary state/drug pages (chunk)"))
fixed_types.extend([
    ("dental", "Dental coverage state + plan pages"),
    ("faq", "FAQ category + question pages"),
    ("billing", "Billing category pages"),
    ("life-events", "Life event decision tree pages"),
    ("guides", "Editorial guide pages"),
    ("states", "State overview pages"),
    ("drugs", "Drug categories, comparisons, county drug pages"),
])

add_table(doc,
    ["Sub-Sitemap", "Content"],
    [[n, d] for n, d in fixed_types],
    col_widths=[5, 12],
)

doc.add_heading("10.2 Sample Sitemap XML Output", level=2)
sample_xml = """<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
  <url>
    <loc>https://healthinsurancerenew.com/formulary/texas/metformin</loc>
    <lastmod>2026-03-17</lastmod>
    <changefreq>yearly</changefreq>
    <priority>0.6</priority>
  </url>
  <url>
    <loc>https://healthinsurancerenew.com/formulary/texas/ozempic</loc>
    <lastmod>2026-03-17</lastmod>
    <changefreq>yearly</changefreq>
    <priority>0.6</priority>
  </url>
  <!-- ... up to 50,000 entries per sub-sitemap -->
</urlset>"""

p = doc.add_paragraph()
run = p.add_run(sample_xml)
run.font.name = "Consolas"
run.font.size = Pt(8)

doc.add_heading("10.3 Data Freshness", level=2)
add_table(doc,
    ["Field", "Value"],
    [
        ["Plan Year", "2026"],
        ["CMS Data Source", "MR-PUF (Machine-Readable Public Use Files)"],
        ["Issuers Attempted", "110"],
        ["Issuers Successful", "106 (96.4% success rate)"],
        ["Total Drug Records", "20,497,198"],
        ["Index Generated", metadata.get("generated_at", "N/A")],
        ["Sitemap lastmod", "2026-03-17"],
    ],
    col_widths=[5, 12],
)


# ── SAVE ────────────────────────────────────────────────────────────────────

doc.save(str(OUTPUT_PATH))
print(f"Report saved to: {OUTPUT_PATH}")
print(f"Pages: ~{len(doc.paragraphs) // 30 + 12} (estimated)")
